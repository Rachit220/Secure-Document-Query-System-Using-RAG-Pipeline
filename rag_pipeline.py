import os
import logging
import hashlib
from typing import List, Dict, Tuple, Optional, Any
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from io import BytesIO
from dotenv import load_dotenv

load_dotenv()

import PyPDF2
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain_core.documents import Document

from langchain_core.embeddings import Embeddings
from langchain_core.language_models.chat_models import BaseChatModel
from langchain_core.outputs import ChatResult, ChatGeneration
from langchain_core.messages import BaseMessage, AIMessage
from sanitizer import PiiSanitizer

logger = logging.getLogger(__name__)


class MockEmbeddings(Embeddings):
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        vectors = []
        for text in texts:
            val = float(hash(text) % 1000) / 1000.0
            vectors.append([val] * 1536)
        return vectors

    def embed_query(self, text: str) -> List[float]:
        val = float(hash(text) % 1000) / 1000.0
        return [val] * 1536


class MockChatOpenAI(BaseChatModel):
    model_name: str = 'mock-gpt-4o-mini'

    def _generate(self, messages: List[BaseMessage], stop: Optional[List[str]] = None, run_manager: Optional[Any] = None, **kwargs: Any) -> ChatResult:
        prompt_text = ''
        for msg in messages:
            prompt_text += msg.content + '\n'

        context = ''
        if 'context:' in prompt_text.lower():
            parts = prompt_text.split('Context:')
            if len(parts) > 1:
                context = parts[1].split('Question:')[0].strip()

        question = ''
        if 'question:' in prompt_text.lower():
            parts = prompt_text.split('Question:')
            if len(parts) > 1:
                question = parts[1].split('\n')[0].strip()

        answer = f"[MOCK DEMO RESPONSE] Based on the retrieved tenant documents, the answer to '{question or 'your question'}' is:\n\n"
        if context:
            lines = [line.strip() for line in context.split('\n') if line.strip() and (not line.startswith('---'))]
            if lines:
                answer += f'Detected relevant content: "{lines[0][:150]}..."'
            else:
                answer += 'Matching records were successfully retrieved.'
        else:
            answer += 'No matching documents were found in the database.'

        ai_message = AIMessage(content=answer)
        return ChatResult(generations=[ChatGeneration(message=ai_message)])

    @property
    def _llm_type(self) -> str:
        return 'mock-chat-openai'


@dataclass
class UploadMetadata:
    tenant_id: str
    filename: str
    file_hash: str
    upload_timestamp: str
    file_size: int
    content_length: int
    chunk_count: int


class DocumentExtractor:
    @staticmethod
    def extract_pdf(file_content: bytes) -> str:
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ''
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += f'\n--- Page {page_num + 1} ---\n'
                text += page.extract_text()
                
            logger.info(f'Extracted {len(pdf_reader.pages)} pages from PDF')
            return text
            
        except Exception as e:
            logger.error(f'PDF extraction failed: {e}')
            raise

    @staticmethod
    def extract_docx(file_content: bytes) -> str:
        try:
            doc = DocxDocument(BytesIO(file_content))
            text = ''
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + '\n'
                    
            logger.info(f'Extracted {len(doc.paragraphs)} paragraphs from DOCX')
            return text
            
        except Exception as e:
            logger.error(f'DOCX extraction failed: {e}')
            raise

    @staticmethod
    def extract(filename: str, file_content: bytes) -> str:
        ext = Path(filename).suffix.lower()
        
        if ext == '.pdf':
            return DocumentExtractor.extract_pdf(file_content)
        elif ext == '.docx':
            return DocumentExtractor.extract_docx(file_content)
        else:
            raise ValueError(f'Unsupported file format: {ext}')


class TenantIsolatedChromaDB:
    def __init__(self, persist_directory: str = './chroma_db', embedding_model: str = 'text-embedding-3-small'):
        self.persist_directory = persist_directory
        self.embedding_model = embedding_model

        if os.getenv('MOCK_AI', 'false').lower() == 'true':
            self.embeddings = MockEmbeddings()
            logger.info('Using local MockEmbeddings')
        else:
            self.embeddings = OpenAIEmbeddings(model=embedding_model)

        Path(persist_directory).mkdir(parents=True, exist_ok=True)
        
        self.client = Chroma(
            persist_directory=persist_directory,
            embedding_function=self.embeddings,
            collection_name='secure_documents'
        )
        logger.info(f'Chroma DB initialized at {persist_directory}')

    def add_documents(self, documents: List[Document], tenant_id: str, source_filename: str) -> int:
        try:
            for doc in documents:
                doc.metadata['tenant_id'] = tenant_id
                doc.metadata['source_file'] = source_filename
                doc.metadata['ingestion_timestamp'] = datetime.utcnow().isoformat()
                
            self.client.add_documents(documents)
            logger.info("Added documents to ChromaDB")
            return len(documents)
            
        except Exception as e:
            logger.error(f'Failed to add documents to ChromaDB: {e}')
            raise

    def get_tenant_retriever(self, tenant_id: str, k: int = 5):
        filter_dict = {'tenant_id': {'$eq': tenant_id}}
        
        retriever = self.client.as_retriever(
            search_type='similarity',
            search_kwargs={'k': k, 'filter': filter_dict}
        )
        logger.debug(f'Created retriever for tenant {tenant_id}')
        return retriever

    def similarity_search(self, query: str, tenant_id: str, k: int = 5) -> List[Tuple[Document, float]]:
        try:
            filter_dict = {'tenant_id': {'$eq': tenant_id}}
            results = self.client.similarity_search_with_score(query, k=k, filter=filter_dict)
            logger.debug("Similarity search completed")
            return results
            
        except Exception as e:
            logger.error(f'Similarity search failed: {e}')
            raise


class RAGPipeline:
    def __init__(
        self,
        chroma_persist_dir: str = './chroma_db',
        embedding_model: str = 'text-embedding-3-small',
        llm_model: str = 'gpt-4o-mini',
        chunk_size: int = 1000,
        chunk_overlap: int = 100
    ):
        self.chroma_persist_dir = chroma_persist_dir
        self.embedding_model = embedding_model
        self.llm_model = llm_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        self.sanitizer = PiiSanitizer(use_presidio=True)
        
        self.vector_db = TenantIsolatedChromaDB(
            persist_directory=chroma_persist_dir,
            embedding_model=embedding_model
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=['\n\n', '\n', ' ', '']
        )
        
        if os.getenv('MOCK_AI', 'false').lower() == 'true':
            self.llm = MockChatOpenAI()
            logger.info('Using local MockChatOpenAI')
        else:
            self.llm = ChatOpenAI(model_name=llm_model, temperature=0)
            
        logger.info('RAG Pipeline initialized successfully')

    def ingest_document(self, tenant_id: str, filename: str, file_content: bytes) -> UploadMetadata:
        logger.info(f'Starting ingestion for tenant {tenant_id}, file {filename}')
        
        raw_text = DocumentExtractor.extract(filename, file_content)
        sanitized_text, entities = self.sanitizer.sanitize(raw_text)
        chunks = self.text_splitter.split_text(sanitized_text)
        
        documents = [
            Document(
                page_content=chunk,
                metadata={'chunk_id': i, 'tenant_id': tenant_id, 'source_file': filename}
            )
            for i, chunk in enumerate(chunks)
        ]
        
        stored_count = self.vector_db.add_documents(documents, tenant_id, filename)
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        metadata = UploadMetadata(
            tenant_id=tenant_id,
            filename=filename,
            file_hash=file_hash,
            upload_timestamp=datetime.utcnow().isoformat(),
            file_size=len(file_content),
            content_length=len(sanitized_text),
            chunk_count=stored_count
        )
        
        logger.info(f'Document ingestion complete: {metadata}')
        return metadata

    def query(self, tenant_id: str, user_query: str, k: int = 5) -> Dict[str, Any]:
        logger.info(f'Query from tenant {tenant_id}: {user_query[:100]}...')
        
        sanitized_query, query_entities = self.sanitizer.sanitize(user_query)
        retrieval_results = self.vector_db.similarity_search(sanitized_query, tenant_id=tenant_id, k=k)
        
        if not retrieval_results:
            return {
                'answer': 'I cannot find relevant information in your uploaded documents.',
                'sources': [],
                'context_chunks': 0,
                'confidence': 0.0
            }

        context_docs = [doc for doc, score in retrieval_results]
        context_text = '\n\n---\n\n'.join([doc.page_content for doc in context_docs])
        scores = [score for doc, score in retrieval_results]

        system_prompt = 'You are an enterprise document assistant. Answer the user\'s question using ONLY the provided context below.\n\nCRITICAL RULES:\n1. Only use information explicitly present in the context.\n2. If the information is not contained in the context, explicitly state: "I cannot find relevant information in your uploaded documents."\n3. DO NOT invent, extrapolate, or assume facts not in the context.\n4. Cite the specific section of the context when answering.\n5. Be concise and professional.'
        
        prompt_template = PromptTemplate(
            input_variables=['context', 'question'],
            template=f'{system_prompt}\n\nCONTEXT FROM YOUR DOCUMENTS:\n{{context}}\n\nUSER QUESTION:\n{{question}}\n\nRESPONSE:'
        )
        
        try:
            formatted_prompt = prompt_template.format(context=context_text, question=sanitized_query)
            response = self.llm.invoke(formatted_prompt)
            answer = response.content
        except Exception as e:
            logger.error(f'LLM invocation failed: {e}')
            answer = 'I encountered an error processing your query. Please try again.'

        sources = [{'filename': doc.metadata.get('source_file', 'unknown'), 'chunk_id': doc.metadata.get('chunk_id', -1)} for doc in context_docs]
        
        result = {
            'answer': answer,
            'sources': sources,
            'context_chunks': len(context_docs),
            'average_relevance_score': sum(scores) / len(scores) if scores else 0.0,
            'tenant_id': tenant_id,
            'sanitization_needed': len(query_entities) > 0
        }
        
        logger.info(f'Query completed successfully for tenant {tenant_id}')
        return result